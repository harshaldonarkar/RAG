# Agentic RAG-Powered AI Research Assistant (Open Source LLMs)
# Built for AllyNerds NLP Engineer Task - No API Keys Required!

import os
import re
import json
import asyncio
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from abc import ABC, abstractmethod
import logging
import requests
import time
import numpy as np
import torch
import faiss
from transformers import pipeline
from sentence_transformers import SentenceTransformer
import shutil
from pathlib import Path
import importlib
from typing import Callable
import google.generativeai as genai 
import groq as groq_mod  
import openai as openai_mod
try:
    import anthropic as anthropic_mod
except ImportError:
    anthropic_mod = None
from bs4 import BeautifulSoup 
from pypdf import PdfReader  
from docx import Document as DocxDocument  
from dotenv import load_dotenv

try:
    from rank_bm25 import BM25Okapi
    _BM25_AVAILABLE = True
except ImportError:
    BM25Okapi = None
    _BM25_AVAILABLE = False


# LangChain / LangGraph (required)
from langchain_community.embeddings import HuggingFaceEmbeddings as LCHuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS as LCFAISS
from langchain_community.llms import Ollama as LCOllama
from langchain_community.llms.huggingface_pipeline import HuggingFacePipeline as LCHuggingFacePipeline
from langgraph.graph import StateGraph, END


# (Dependency presence is checked in __main__ using importlib)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ----------------------------
# Constants and configuration
# ----------------------------
# Note: chunking is word-based, not token-based
DEFAULT_CHUNK_SIZE_WORDS: int = 200
DEFAULT_CHUNK_OVERLAP_WORDS: int = 40
DEFAULT_TOP_K: int = 3
SUPPORTED_FILE_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

@dataclass
class ResearchQuery:
    """Data class for research queries"""
    query: str
    intent: str = ""
    sub_queries: List[str] = field(default_factory=list)
    priority: int = 1

@dataclass
class ResearchResult:
    """Data class for research results"""
    content: str
    sources: List[str]
    confidence: float
    agent_chain: List[str]

@dataclass
class Document:
    """Simple document class"""
    content: str
    metadata: Dict[str, Any]

class LLMProvider:
    """Manages different LLM backends (local models, Ollama, Hugging Face)"""
    
    def __init__(self, provider_type: str = "huggingface", model_name: Optional[str] = None):
        self.provider_type = provider_type
        # Allow override via env var, fallback to a lightweight default
        self.model_name = model_name or os.getenv("HF_MODEL_NAME", "distilgpt2")
        self.model = None
        self.tokenizer = None
        self.pipeline = None
        self._initialize_model()
    
    def _initialize_model(self):
        """Initialize the chosen LLM backend"""
        logger.info(f"Initializing {self.provider_type} model: {self.model_name}")
        
        if self.provider_type == "huggingface":
            try:
                # Choose pipeline type based on model family (T5/FLAN → text2text)
                model_name_lower = self.model_name.lower()
                if any(keyword in model_name_lower for keyword in ["t5", "flan"]):
                    task = "text2text-generation"
                else:
                    task = "text-generation"

                self.pipeline = pipeline(
                    task,
                    model=self.model_name,
                    device=0 if torch.cuda.is_available() else -1,
                )
                logger.info(f"Hugging Face model loaded: task={task}, model={self.model_name}")
            except Exception as e:
                logger.error(f"Error loading HF model '{self.model_name}': {e}")
                # Fallback to even simpler approach
                self._initialize_fallback()
        
        elif self.provider_type == "ollama":
            # Ollama integration (requires Ollama to be installed and running)
            self.ollama_url = "http://localhost:11434/api/generate"
            # Prefer explicitly provided model_name, else env var, else sensible default
            self.ollama_model_name = self.model_name or os.getenv("OLLAMA_MODEL", "mistral:7b-instruct-q5_K_M")
            # Options via env vars
            self.ollama_num_ctx = int(os.getenv("OLLAMA_NUM_CTX", "8192"))
            self.ollama_temperature = float(os.getenv("OLLAMA_TEMPERATURE", "0.3"))
            logger.info(f"Using Ollama model: {self.ollama_model_name} (num_ctx={self.ollama_num_ctx}, temp={self.ollama_temperature})")
            logger.info("Make sure Ollama is running: ollama serve")
        
        elif self.provider_type == "gemini":
            try:
                if genai is None:
                    raise RuntimeError("google-generativeai package not installed")
                api_key = os.getenv("GOOGLE_API_KEY")
                if not api_key:
                    raise RuntimeError("GOOGLE_API_KEY is not set")
                genai.configure(api_key=api_key)
                self.gemini_model_name = self.model_name or os.getenv("GEMINI_MODEL", "gemini-2.0-flash")
                self.gemini_client = genai.GenerativeModel(self.gemini_model_name)
                logger.info(f"Google Gemini model ready: {self.gemini_model_name}")
            except Exception as e:
                logger.error(f"Error initializing Gemini: {e}")
                self._initialize_fallback()
        
        elif self.provider_type == "groq":
            try:
                if groq_mod is None:
                    raise RuntimeError("groq package not installed")
                api_key = os.getenv("GROQ_API_KEY")
                if not api_key:
                    raise RuntimeError("GROQ_API_KEY is not set")
                self.groq_client = groq_mod.Groq(api_key=api_key)
                self.groq_model_name = self.model_name or os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
                logger.info(f"Groq client ready: {self.groq_model_name}")
            except Exception as e:
                logger.error(f"Error initializing Groq: {e}")
                self._initialize_fallback()

        elif self.provider_type == "openai":
            try:
                if openai_mod is None:
                    raise RuntimeError("openai package not installed")
                api_key = os.getenv("OPENAI_API_KEY")
                if not api_key:
                    raise RuntimeError("OPENAI_API_KEY is not set")
                self.openai_client = openai_mod.OpenAI(api_key=api_key)
                self.openai_model_name = self.model_name or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
                self.openai_temperature = float(os.getenv("OPENAI_TEMPERATURE", "0.3"))
                logger.info(f"OpenAI client ready: {self.openai_model_name}")
            except Exception as e:
                logger.error(f"Error initializing OpenAI: {e}")
                self._initialize_fallback()

        elif self.provider_type == "claude":
            try:
                if anthropic_mod is None:
                    raise RuntimeError("anthropic package not installed")
                api_key = os.getenv("ANTHROPIC_API_KEY")
                if not api_key:
                    raise RuntimeError("ANTHROPIC_API_KEY is not set")
                self.claude_client = anthropic_mod.Anthropic(api_key=api_key)
                self.claude_model_name = self.model_name or os.getenv("CLAUDE_MODEL", "claude-sonnet-4-6")
                self.claude_temperature = float(os.getenv("CLAUDE_TEMPERATURE", "0.3"))
                logger.info(f"Claude client ready: {self.claude_model_name}")
            except Exception as e:
                logger.error(f"Error initializing Claude: {e}")
                self._initialize_fallback()

        else:
            self._initialize_fallback()
    
    def _initialize_fallback(self):
        """Fallback to rule-based responses when models fail"""
        self.provider_type = "fallback"
        logger.warning("Using fallback rule-based responses")
    
    async def generate_text(self, prompt: str, max_tokens: int = 200) -> str:
        """Generate text using the configured LLM"""
        
        if self.provider_type == "huggingface" and self.pipeline:
            try:
                # Generate with the pipeline
                result = self.pipeline(
                    prompt,
                    max_new_tokens=max_tokens,
                    temperature=0.7,
                    do_sample=True,
                    pad_token_id=getattr(self.pipeline.tokenizer, 'eos_token_id', None)
                )
                generated_text = result[0].get('generated_text', '')
                # For text-generation, strip the prompt; for text2text, return as-is
                task = getattr(self.pipeline, 'task', '')
                if task == 'text-generation' and generated_text.startswith(prompt):
                    new_text = generated_text[len(prompt):].strip()
                else:
                    new_text = generated_text.strip()
                return new_text if new_text else "I need more context to provide a detailed response."
            except Exception as e:
                logger.error(f"HuggingFace generation error: {e}")
                return self._fallback_response(prompt)
        
        elif self.provider_type == "ollama":
            return await self._generate_with_ollama(prompt)
        
        elif self.provider_type == "gemini":
            return await self._generate_with_gemini(prompt, max_tokens)
        
        elif self.provider_type == "groq":
            return await self._generate_with_groq(prompt, max_tokens)
        
        elif self.provider_type == "openai":
            return await self._generate_with_openai(prompt, max_tokens)

        elif self.provider_type == "claude":
            return await self._generate_with_claude(prompt, max_tokens)

        else:
            return self._fallback_response(prompt)
    
    async def _generate_with_ollama(self, prompt: str) -> str:
        """Generate text using local Ollama instance"""
        try:
            data = {
                "model": getattr(self, "ollama_model_name", os.getenv("OLLAMA_MODEL", "mistral:7b-instruct-q5_K_M")),
                "prompt": prompt,
                "stream": False,
                "options": {
                    "num_ctx": getattr(self, "ollama_num_ctx", int(os.getenv("OLLAMA_NUM_CTX", "8192"))),
                    "temperature": getattr(self, "ollama_temperature", float(os.getenv("OLLAMA_TEMPERATURE", "0.3")))
                }
            }
            
            response = requests.post(self.ollama_url, json=data, timeout=30)
            if response.status_code == 200:
                return response.json().get("response", "No response generated")
            else:
                return self._fallback_response(prompt)
        
        except Exception as e:
            logger.error(f"Ollama error: {e}")
            return self._fallback_response(prompt)

    async def _generate_with_gemini(self, prompt: str, max_tokens: int) -> str:
        try:
            loop = asyncio.get_event_loop()
            def _call():
                generation_config = {"max_output_tokens": max_tokens, "temperature": 0.7}
                response = self.gemini_client.generate_content(prompt, generation_config=generation_config)
                try:
                    return response.text or ""
                except Exception:
                    # Fallback extraction
                    if hasattr(response, "candidates") and response.candidates:
                        parts = getattr(response.candidates[0], "content", None)
                        if parts and hasattr(parts, "parts") and parts.parts:
                            return getattr(parts.parts[0], "text", "")
                    return ""
            text = await loop.run_in_executor(None, _call)
            return text if text else self._fallback_response(prompt)
        except Exception as e:
            logger.error(f"Gemini error: {e}")
            return self._fallback_response(prompt)

    async def _generate_with_groq(self, prompt: str, max_tokens: int) -> str:
        try:
            loop = asyncio.get_event_loop()
            def _call():
                completion = self.groq_client.chat.completions.create(
                    model=self.groq_model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=max_tokens,
                )
                return completion.choices[0].message.content
            text = await loop.run_in_executor(None, _call)
            return text if text else self._fallback_response(prompt)
        except Exception as e:
            logger.error(f"Groq error: {e}")
            return self._fallback_response(prompt)

    async def _generate_with_openai(self, prompt: str, max_tokens: int) -> str:
        try:
            loop = asyncio.get_event_loop()
            def _call():
                completion = self.openai_client.chat.completions.create(
                    model=self.openai_model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=self.openai_temperature,
                    max_tokens=max_tokens,
                )
                return completion.choices[0].message.content
            text = await loop.run_in_executor(None, _call)
            return text if text else self._fallback_response(prompt)
        except Exception as e:
            logger.error(f"OpenAI error: {e}")
            return self._fallback_response(prompt)
    
    async def _generate_with_claude(self, prompt: str, max_tokens: int) -> str:
        try:
            loop = asyncio.get_event_loop()
            def _call():
                message = self.claude_client.messages.create(
                    model=self.claude_model_name,
                    max_tokens=max_tokens,
                    temperature=self.claude_temperature,
                    messages=[{"role": "user", "content": prompt}],
                )
                return message.content[0].text
            text = await loop.run_in_executor(None, _call)
            return text if text else self._fallback_response(prompt)
        except Exception as e:
            logger.error(f"Claude error: {e}")
            return self._fallback_response(prompt)

    def _fallback_response(self, prompt: str) -> str:
        """Rule-based fallback responses"""
        prompt_lower = prompt.lower()
        
        if "query" in prompt_lower and "intent" in prompt_lower:
            return json.dumps({
                "intent": "Research and information gathering",
                "sub_queries": [prompt.split(":")[-1].strip() if ":" in prompt else prompt],
                "priority": 3,
                "research_domain": "General"
            })
        
        elif "plan" in prompt_lower and ("research" in prompt_lower or "steps" in prompt_lower):
            return json.dumps({
                "steps": ["Gather relevant information", "Analyze key concepts", "Synthesize findings"],
                "focus_areas": ["Core concepts", "Recent developments", "Practical applications"],
                "output_structure": ["Introduction", "Main Findings", "Key Insights", "Conclusion"],
                "complexity": "medium"
            })
        
        elif "summary" in prompt_lower or "synthesize" in prompt_lower:
            return """Based on the available information, here are the key findings:

1. **Overview**: The research topic involves multiple interconnected concepts that require careful analysis.

2. **Key Findings**: 
   - Current developments show significant progress in the field
   - Multiple approaches are being explored by researchers
   - Practical applications are emerging across various domains

3. **Implications**: These findings suggest continued growth and innovation in this area.

4. **Conclusion**: Further research and development will likely yield additional insights and applications."""
        
        else:
            return f"Based on the query about '{prompt[:50]}...', this appears to be a research topic that would benefit from systematic investigation of current literature and recent developments in the field."

class EmbeddingProvider:
    """Handles document embeddings using sentence-transformers"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model_name = model_name
        try:
            self.model = SentenceTransformer(model_name)
            logger.info(f"✅ Embedding model '{model_name}' loaded successfully")
        except Exception as e:
            logger.error(f"Error loading embedding model: {e}")
            self.model = None
    
    def encode(self, texts: List[str]) -> np.ndarray:
        """Encode texts to embeddings, falling back to random demo embeddings if needed."""
        if self.model:
            return self.model.encode(texts)
        # Fallback: random embeddings (demo only)
        logger.warning("Using random embeddings - install sentence-transformers for real embeddings")
        return np.random.randn(len(texts), 384).astype(np.float32)

class Reranker:
    """Cross-encoder reranker for re-scoring retrieved candidates.

    Uses a small MS-MARCO cross-encoder by default — scores each (query, passage)
    pair jointly, which is far more accurate than embedding cosine similarity alone.
    """

    DEFAULT_MODEL = "cross-encoder/ms-marco-MiniLM-L-6-v2"

    def __init__(self, model_name: str = DEFAULT_MODEL):
        self.model_name = model_name
        self.model = None
        self._load()

    def _load(self) -> None:
        try:
            from sentence_transformers.cross_encoder import CrossEncoder as CE
            self.model = CE(self.model_name)
            logger.info(f"✅ Reranker loaded: {self.model_name}")
        except Exception as e:
            logger.warning(f"Reranker unavailable ({e}); skipping reranking.")
            self.model = None

    def rerank(self, query: str, documents: List[Document], top_k: int) -> List[Document]:
        """Return top_k documents sorted by cross-encoder relevance score."""
        if self.model is None or not documents:
            return documents[:top_k]
        try:
            pairs = [(query, doc.content) for doc in documents]
            scores = self.model.predict(pairs)
            scored = sorted(zip(scores, documents), key=lambda x: x[0], reverse=True)
            results = []
            for score, doc in scored[:top_k]:
                meta = dict(doc.metadata)
                meta["rerank_score"] = float(score)
                results.append(Document(content=doc.content, metadata=meta))
            return results
        except Exception as e:
            logger.error(f"Reranking error: {e}")
            return documents[:top_k]


class WebSearcher:
    """Fetches live web results to supplement local RAG retrieval.

    Priority:
      1. Tavily  — if TAVILY_API_KEY is set (clean, AI-optimised excerpts)
      2. DuckDuckGo — free fallback, no API key required
    """

    def __init__(self) -> None:
        self._tavily = None
        self._ddg_available = False
        self._initialize()

    def _initialize(self) -> None:
        tavily_key = os.getenv("TAVILY_API_KEY")
        if tavily_key:
            try:
                from tavily import TavilyClient
                self._tavily = TavilyClient(api_key=tavily_key)
                logger.info("✅ Tavily web search ready")
                return
            except Exception as e:
                logger.warning(f"Tavily init failed: {e}")
        try:
            from duckduckgo_search import DDGS  # noqa: F401
            self._ddg_available = True
            logger.info("✅ DuckDuckGo web search ready (no API key)")
        except ImportError:
            logger.warning("Web search unavailable — install duckduckgo-search or set TAVILY_API_KEY")

    @property
    def is_available(self) -> bool:
        return self._tavily is not None or self._ddg_available

    async def search(self, query: str, max_results: int = 4) -> List[Document]:
        """Return web results as Document objects, or [] if unavailable."""
        if self._tavily:
            return await self._tavily_search(query, max_results)
        if self._ddg_available:
            return await self._ddg_search(query, max_results)
        return []

    async def _tavily_search(self, query: str, max_results: int) -> List[Document]:
        loop = asyncio.get_event_loop()
        def _call():
            resp = self._tavily.search(query, max_results=max_results, search_depth="basic")
            docs = []
            for r in resp.get("results", []):
                content = f"{r.get('title', '')}\n\n{r.get('content', '')}"
                docs.append(Document(
                    content=content,
                    metadata={"source": r.get("url", "web"), "type": "web", "score": r.get("score", 0.5)},
                ))
            return docs
        try:
            return await loop.run_in_executor(None, _call)
        except Exception as e:
            logger.error(f"Tavily search error: {e}")
            return []

    async def _ddg_search(self, query: str, max_results: int) -> List[Document]:
        loop = asyncio.get_event_loop()
        def _call():
            from duckduckgo_search import DDGS
            docs = []
            with DDGS() as ddgs:
                for r in ddgs.text(query, max_results=max_results):
                    content = f"{r.get('title', '')}\n\n{r.get('body', '')}"
                    docs.append(Document(
                        content=content,
                        metadata={"source": r.get("href", "web"), "type": "web", "score": 0.5},
                    ))
            return docs
        try:
            return await loop.run_in_executor(None, _call)
        except Exception as e:
            logger.error(f"DuckDuckGo search error: {e}")
            return []


class ConversationMemory:
    """Stores recent conversation turns and provides context for follow-up queries.

    Two uses:
      1. Query rewriting — resolve references like "it", "that", "the second point"
         into standalone queries before hitting the retrieval pipeline.
      2. Synthesis context — let the synthesizer know what was already covered so
         the new report builds on (not repeats) prior answers.
    """

    MAX_TURNS = 6  # keep last 6 turns (3 exchanges)

    def __init__(self) -> None:
        self.turns: List[Dict[str, str]] = []  # [{"query": ..., "summary": ...}]

    def add(self, query: str, summary: str) -> None:
        self.turns.append({"query": query, "summary": summary[:400]})
        if len(self.turns) > self.MAX_TURNS:
            self.turns.pop(0)

    def is_empty(self) -> bool:
        return len(self.turns) == 0

    def format_for_rewrite(self) -> str:
        """Compact history for the query-rewriting prompt."""
        return "\n".join(
            f"Turn {i}: User asked: \"{t['query']}\""
            for i, t in enumerate(self.turns, 1)
        )

    def format_for_synthesis(self) -> str:
        """Prior Q&A pairs injected into the synthesis prompt."""
        lines = []
        for t in self.turns:
            lines.append(f"Previous question: {t['query']}\nPrevious answer summary: {t['summary']}")
        return "\n\n".join(lines)


class Agent(ABC):
    """Abstract base class for all agents"""
    
    def __init__(self, name: str, llm_provider: LLMProvider):
        self.name = name
        self.llm_provider = llm_provider
        self.context = {}
    
    @abstractmethod
    async def process(self, input_data: Any) -> Any:
        """Process input and return output"""
        pass

class QueryUnderstandingAgent(Agent):
    """Agent that analyzes and breaks down user queries"""

    def __init__(self, llm_provider: LLMProvider):
        super().__init__("QueryUnderstandingAgent", llm_provider)

    async def rewrite_query(self, query: str, memory: "ConversationMemory") -> str:
        """Rewrite a follow-up query into a self-contained standalone query.

        Resolves references like "tell me more", "what about the second point",
        "and safety?" using the conversation history.
        """
        if memory.is_empty():
            return query
        prompt = (
            f"Given this conversation history:\n{memory.format_for_rewrite()}\n\n"
            f"Rewrite the following follow-up message as a complete, standalone research query "
            f"that can be understood without the history. "
            f"If it is already self-contained, return it unchanged. "
            f"Output only the rewritten query, nothing else.\n\n"
            f"Follow-up: \"{query}\"\nStandalone query:"
        )
        try:
            rewritten = await self.llm_provider.generate_text(prompt, max_tokens=80)
            rewritten = rewritten.strip().strip('"').strip("'")
            return rewritten if len(rewritten) > 5 else query
        except Exception:
            return query

    async def process(self, query: str) -> ResearchQuery:
        """Analyze query intent and break it down into sub-queries"""
        
        prompt = f"""Analyze this research query and provide JSON output:
Query: "{query}"

Provide analysis in JSON format:
{{
    "intent": "brief description of what user wants",
    "sub_queries": ["focused sub-query 1", "focused sub-query 2"],
    "priority": 3,
    "research_domain": "domain like AI, technology, science, etc."
}}

JSON:"""
        
        response = await self.llm_provider.generate_text(prompt, max_tokens=200)
        
        try:
            # Try to extract JSON from response
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response[json_start:json_end]
                analysis = json.loads(json_str)
                return ResearchQuery(
                    query=query,
                    intent=analysis.get("intent", "Research query"),
                    sub_queries=analysis.get("sub_queries", [query]),
                    priority=analysis.get("priority", 3)
                )
        except:
            pass
        
        # Fallback: simple rule-based analysis
        sub_queries = []
        if " and " in query.lower():
            sub_queries = [q.strip() for q in query.lower().split(" and ")]
        elif " in " in query.lower():
            sub_queries = [query, query.split(" in ")[0].strip()]
        else:
            sub_queries = [query]
        
        return ResearchQuery(
            query=query,
            intent=f"Research about {query}",
            sub_queries=sub_queries[:3],  # Max 3 sub-queries
            priority=3
        )

class SemanticChunker:
    """Splits documents at topic-shift boundaries detected via embedding similarity.

    Algorithm:
      1. Split text into sentences (regex, no extra deps)
      2. Batch-embed all sentences
      3. Walk consecutive pairs: when cosine similarity drops below `split_threshold`
         OR the running chunk would exceed `max_words`, close the current chunk
      4. Merge orphan sentences (<= MIN_SENTENCE_WORDS) into the previous chunk
      5. Falls back to the caller's word-based chunker on any error

    Produces coherent chunks where each chunk stays on one topic.
    """

    MIN_SENTENCE_WORDS = 4   # discard/merge sentences shorter than this
    SPLIT_THRESHOLD    = 0.35  # cosine similarity below which we start a new chunk

    # Regex that splits on sentence-ending punctuation followed by whitespace
    _SENTENCE_RE = re.compile(r'(?<=[.!?])\s+')

    def __init__(self, embedding_provider: "EmbeddingProvider") -> None:
        self.embedding_provider = embedding_provider

    def chunk(self, text: str, max_words: int = 200) -> List[str]:
        """Return a list of semantically coherent text chunks."""
        sentences = self._split_sentences(text)
        if len(sentences) <= 2:
            # Too short to bother with semantics — return as one chunk
            return [text.strip()] if text.strip() else []

        try:
            return self._semantic_split(sentences, max_words)
        except Exception as e:
            logger.warning(f"Semantic chunking failed ({e}); falling back to sentence groups")
            return self._naive_sentence_groups(sentences, max_words)

    # ------------------------------------------------------------------
    def _split_sentences(self, text: str) -> List[str]:
        raw = self._SENTENCE_RE.split(text.strip())
        return [s.strip() for s in raw if len(s.split()) >= self.MIN_SENTENCE_WORDS]

    def _semantic_split(self, sentences: List[str], max_words: int) -> List[str]:
        # Batch embed (reuse existing provider — no extra model load)
        embs = self.embedding_provider.encode(sentences).astype(np.float32)
        # Normalise for cosine similarity
        norms = np.linalg.norm(embs, axis=1, keepdims=True)
        norms = np.where(norms == 0.0, 1.0, norms)
        embs = embs / norms

        chunks: List[str] = []
        current: List[str] = [sentences[0]]
        current_words: int = len(sentences[0].split())

        for i in range(1, len(sentences)):
            sim = float(np.dot(embs[i - 1], embs[i]))
            sw = len(sentences[i].split())
            topic_shift = sim < self.SPLIT_THRESHOLD
            size_overflow = (current_words + sw) > max_words

            if (topic_shift or size_overflow) and current:
                chunks.append(" ".join(current))
                current = [sentences[i]]
                current_words = sw
            else:
                current.append(sentences[i])
                current_words += sw

        if current:
            chunks.append(" ".join(current))

        # Drop empty / tiny chunks
        return [c for c in chunks if len(c.split()) >= self.MIN_SENTENCE_WORDS]

    def _naive_sentence_groups(self, sentences: List[str], max_words: int) -> List[str]:
        """Simple fallback: group sentences until max_words is hit."""
        chunks, current, current_words = [], [], 0
        for s in sentences:
            sw = len(s.split())
            if current and current_words + sw > max_words:
                chunks.append(" ".join(current))
                current, current_words = [], 0
            current.append(s)
            current_words += sw
        if current:
            chunks.append(" ".join(current))
        return chunks


class RAGRetriever:
    """RAG pipeline using sentence-transformers and FAISS (cosine similarity, overlapping chunks)."""
    
    def __init__(
        self,
        documents_path: str = "./documents",
        chunk_size_words: int = DEFAULT_CHUNK_SIZE_WORDS,
        chunk_overlap_words: int = DEFAULT_CHUNK_OVERLAP_WORDS,
        use_cosine_similarity: bool = True,
        embedding_model_name: str = "all-MiniLM-L6-v2",
        use_hybrid: bool = True,
        use_reranking: bool = True,
        reranker_model: str = Reranker.DEFAULT_MODEL,
        use_semantic_chunking: bool = True,
    ):
        self.documents_path = documents_path
        self.embedding_provider = EmbeddingProvider(embedding_model_name)
        self.chunk_size_words = max(50, chunk_size_words)
        self.chunk_overlap_words = max(0, min(self.chunk_size_words - 10, chunk_overlap_words))
        self.use_cosine_similarity = use_cosine_similarity
        self.use_hybrid = use_hybrid and _BM25_AVAILABLE
        self.reranker = Reranker(reranker_model) if use_reranking else None
        self.semantic_chunker = SemanticChunker(self.embedding_provider) if use_semantic_chunking else None

        self.documents: List[Document] = []
        self.embeddings: Optional[np.ndarray] = None
        self.index = None
        self.bm25_index = None
        self._id_to_doc_idx: List[int] = []
        self._load_and_index_documents()
    
    def _load_and_index_documents(self):
        """Load documents and create FAISS index"""
        if not os.path.exists(self.documents_path):
            os.makedirs(self.documents_path)
            self._create_sample_documents()
        
        # Load supported files recursively
        self.documents = []
        self._id_to_doc_idx = []
        for root, _, files in os.walk(self.documents_path):
            for filename in files:
                ext = Path(filename).suffix.lower()
                if ext in SUPPORTED_FILE_EXTENSIONS:
                    filepath = os.path.join(root, filename)
                    try:
                        content = self._read_document_file(filepath)
                        if content:
                            # Split using semantic chunking if available, else overlapping word chunks
                            if self.semantic_chunker:
                                chunks = self.semantic_chunker.chunk(content, self.chunk_size_words)
                            else:
                                chunks = self._split_text_into_overlapping_chunks(content)
                            for i, chunk in enumerate(chunks):
                                rel_source = os.path.relpath(filepath, self.documents_path)
                                self.documents.append(Document(
                                    content=chunk,
                                    metadata={'source': rel_source, 'chunk_id': i}
                                ))
                    except Exception as e:
                        logger.error(f"Error reading {filename}: {e}")
        
        if not self.documents:
            self._create_sample_documents()
            self._load_and_index_documents()
            return
        
        # Try loading cached index if corpus unchanged
        try:
            if self._load_index_cache():
                logger.info("Loaded cached FAISS index")
                return
        except Exception as e:
            logger.warning(f"Index cache unavailable, rebuilding: {e}")

        # Create embeddings and FAISS index
        try:
            texts = [doc.content for doc in self.documents]
            self.embeddings = self.embedding_provider.encode(texts).astype(np.float32)
            if self.use_cosine_similarity:
                # Normalize for cosine similarity and use inner-product index
                try:
                    # Use faiss native normalization for speed
                    faiss.normalize_L2(self.embeddings)
                except Exception:
                    self.embeddings = self._normalize_vectors(self.embeddings)
                dimension = self.embeddings.shape[1]
                self.index = faiss.IndexFlatIP(dimension)
            else:
                dimension = self.embeddings.shape[1]
                self.index = faiss.IndexFlatL2(dimension)
            
            self.index.add(self.embeddings)

            logger.info(f"✅ Indexed {len(self.documents)} document chunks")
            # Save cache
            try:
                self._save_index_cache()
            except Exception as e:
                logger.warning(f"Could not save index cache: {e}")

        except Exception as e:
            logger.error(f"Error creating FAISS index: {e}")
            self.index = None

        # Build BM25 index (always rebuilt from in-memory chunks — fast, no cache needed)
        self._build_bm25_index()

    def _read_document_file(self, filepath: str) -> str:
        """Best-effort read for multiple formats: txt, md, html, pdf, docx."""
        ext = Path(filepath).suffix.lower()
        try:
            if ext in {'.txt', '.md'}:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            if ext in {'.html', '.htm'}:
                try:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        html = f.read()
                    if BeautifulSoup:
                        soup = BeautifulSoup(html, 'html.parser')
                        return soup.get_text(separator=' ')
                    # Fallback: naive tag stripping
                    return html.replace('<br>', ' ').replace('<br/>', ' ').replace('<p>', ' ').replace('</p>', ' ')
                except Exception as e:
                    logger.error(f"HTML read error for {filepath}: {e}")
                    return ""
            if ext == '.pdf':
                try:
                    if PdfReader is None:
                        raise RuntimeError("pypdf not installed")
                    reader = PdfReader(filepath)
                    texts = []
                    for page in reader.pages:
                        page_text = page.extract_text() or ''
                        texts.append(page_text)
                    return "\n".join(texts)
                except Exception as e:
                    logger.error(f"PDF read error for {filepath}: {e}")
                    return ""
            if ext == '.docx':
                try:
                    if DocxDocument is None:
                        raise RuntimeError("python-docx not installed")
                    doc = DocxDocument(filepath)
                    return "\n".join(p.text for p in doc.paragraphs)
                except Exception as e:
                    logger.error(f"DOCX read error for {filepath}: {e}")
                    return ""
        except Exception as e:
            logger.error(f"Read error for {filepath}: {e}")
        return ""
    
    def _split_text_into_overlapping_chunks(self, text: str) -> List[str]:
        """Word-based overlapping chunking.
        This avoids breaking mid-sentence while staying simple without extra deps.
        """
        words = text.split()
        if not words:
            return []
        chunks: List[str] = []
        step = max(1, self.chunk_size_words - self.chunk_overlap_words)
        for start in range(0, len(words), step):
            end = min(len(words), start + self.chunk_size_words)
            chunk_words = words[start:end]
            if not chunk_words:
                continue
            chunk_text = " ".join(chunk_words)
            chunks.append(chunk_text)
            if end >= len(words):
                break
        return chunks

    # Backwards-compatible alias used by LangGraph initialization
    def _split_text(self, text: str) -> List[str]:
        return self._split_text_into_overlapping_chunks(text)

    @staticmethod
    def _normalize_vectors(vectors: np.ndarray) -> np.ndarray:
        """L2-normalize vectors row-wise with numerical stability."""
        norms = np.linalg.norm(vectors, axis=1, keepdims=True)
        norms = np.where(norms == 0.0, 1.0, norms)
        return vectors / norms

    # ------------------
    # BM25 / hybrid helpers
    # ------------------
    def _build_bm25_index(self) -> None:
        """Build an in-memory BM25 index from the current document chunks."""
        if not self.use_hybrid or not self.documents:
            self.bm25_index = None
            return
        try:
            tokenized = [doc.content.lower().split() for doc in self.documents]
            self.bm25_index = BM25Okapi(tokenized)
            logger.info(f"✅ BM25 index built over {len(self.documents)} chunks")
        except Exception as e:
            logger.warning(f"BM25 index build failed: {e}")
            self.bm25_index = None

    @staticmethod
    def _reciprocal_rank_fusion(
        dense_indices: List[int],
        sparse_indices: List[int],
        k: int = 60,
    ) -> List[int]:
        """Combine two ranked lists using Reciprocal Rank Fusion (RRF).

        Returns indices sorted by descending fused score.
        k=60 is the standard RRF constant that reduces the impact of very high ranks.
        """
        scores: Dict[int, float] = {}
        for rank, idx in enumerate(dense_indices):
            scores[idx] = scores.get(idx, 0.0) + 1.0 / (k + rank + 1)
        for rank, idx in enumerate(sparse_indices):
            scores[idx] = scores.get(idx, 0.0) + 1.0 / (k + rank + 1)
        return sorted(scores, key=lambda i: scores[i], reverse=True)
    
    def _create_sample_documents(self):
        """Create sample documents for demonstration"""
        sample_docs = {
            "ai_safety.txt": """
AI Safety Research Overview. AI safety is a crucial field focused on ensuring artificial intelligence systems behave safely and beneficially. Key areas include alignment, robustness, and interpretability. Recent trends show increased focus on large language model safety, including work on constitutional AI, reinforcement learning from human feedback (RLHF), and adversarial testing. Major organizations like Anthropic, OpenAI, and DeepMind are leading safety research. Constitutional AI involves training models to follow a set of principles. RLHF uses human feedback to improve model behavior. Adversarial testing helps identify potential failure modes. Research also focuses on value alignment and ensuring AI systems pursue intended goals.
""",
            "machine_learning_trends.txt": """
Machine Learning Trends 2024. Current trends in machine learning include several key areas. Large Language Models (LLMs) and their applications continue to expand rapidly. Multimodal AI combining text, image, and audio processing is becoming more sophisticated. Federated learning enables privacy-preserving machine learning across distributed devices. AutoML and neural architecture search automate model design and optimization. Edge AI and model compression techniques make AI more accessible on mobile devices. Explainable AI (XAI) provides better interpretability of model decisions. Transfer learning allows models to adapt to new tasks with minimal data. Few-shot and zero-shot learning capabilities are improving significantly.
""",
            "nlp_advances.txt": """
Natural Language Processing Advances. Recent advances in NLP have transformed the field significantly. Transformer architectures and attention mechanisms form the backbone of modern NLP systems. Pre-trained language models like BERT, GPT, and T5 have achieved remarkable performance across various tasks. Few-shot and zero-shot learning capabilities allow models to handle new tasks with minimal examples. Multilingual and cross-lingual models enable processing of multiple languages simultaneously. Applications span chatbots, machine translation, text generation, and document analysis. Retrieval-augmented generation (RAG) combines knowledge retrieval with text generation. Fine-tuning techniques adapt pre-trained models to specific domains and tasks.
""",
            "deep_learning.txt": """
Deep Learning Fundamentals. Deep learning uses neural networks with multiple layers to learn complex patterns in data. Convolutional Neural Networks (CNNs) excel at image processing and computer vision tasks. Recurrent Neural Networks (RNNs) and LSTMs handle sequential data like time series and text. Generative Adversarial Networks (GANs) create realistic synthetic data. Variational Autoencoders (VAEs) learn compressed representations of data. Optimization techniques like Adam and RMSprop improve training efficiency. Regularization methods prevent overfitting and improve generalization. GPU computing enables training of large-scale models. Batch normalization and dropout are common techniques for stable training.
""",
            "research_methodology.txt": """
Research Methodology in AI. Systematic research methodology is crucial for advancing AI. Literature review identifies existing knowledge and research gaps. Hypothesis formation guides experimental design and investigation. Data collection and preprocessing ensure high-quality training datasets. Experimental design includes proper controls and statistical analysis. Evaluation metrics measure model performance objectively. Cross-validation prevents overfitting to specific datasets. Baseline comparisons establish performance improvements. Reproducibility requires detailed documentation and code sharing. Peer review ensures quality and validity of research findings. Ethical considerations guide responsible AI research practices.
"""
        }
        
        for filename, content in sample_docs.items():
            filepath = os.path.join(self.documents_path, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content.strip())
        
        logger.info(f"Created {len(sample_docs)} sample documents")
    
    async def retrieve_context(
        self,
        query: str,
        k: int = DEFAULT_TOP_K,
        hyde_passage: Optional[str] = None,
    ) -> List[Document]:
        """Retrieve relevant document chunks using hybrid search + optional reranking.

        Pipeline:
          1. Dense (FAISS): embed hyde_passage (if provided) or raw query
             Sparse (BM25): always uses raw query tokens
             → fuse via RRF → fetch_k candidates
          2. Cross-encoder reranking of candidates → top k
        """
        if not self.index or not self.documents:
            return []

        try:
            # Fetch more candidates when reranking so the cross-encoder has a larger pool
            reranking_on = self.reranker is not None and self.reranker.model is not None
            fetch = min(k * 5 if reranking_on else k * 4, len(self.documents))

            # --- Dense retrieval (FAISS) ---
            # HyDE: use the hypothetical passage for embedding when available;
            # it sits closer to real documents in the embedding space than a short question
            dense_query = hyde_passage if hyde_passage else query
            query_embedding = self.embedding_provider.encode([dense_query]).astype(np.float32)
            if self.use_cosine_similarity:
                try:
                    faiss.normalize_L2(query_embedding)
                except Exception:
                    query_embedding = self._normalize_vectors(query_embedding)

            distances, faiss_indices = self.index.search(query_embedding, fetch)
            dense_ranked: List[int] = [
                int(idx) for idx in faiss_indices[0]
                if 0 <= int(idx) < len(self.documents)
            ]

            # --- Sparse retrieval (BM25) ---
            if self.use_hybrid and self.bm25_index is not None:
                tokens = query.lower().split()
                bm25_scores = self.bm25_index.get_scores(tokens)
                sparse_ranked: List[int] = [
                    int(i) for i in np.argsort(bm25_scores)[::-1][:fetch]
                ]
                fused = self._reciprocal_rank_fusion(dense_ranked, sparse_ranked)
                ranked_indices = fused
                retrieval_mode = "hybrid"
            else:
                ranked_indices = dense_ranked
                retrieval_mode = "dense"

            # --- Build de-duplicated candidate list (up to fetch size) ---
            candidates: List[Document] = []
            seen_keys: set = set()
            for idx in ranked_indices:
                if len(candidates) >= fetch:
                    break
                doc = self.documents[idx]
                key = (doc.metadata.get("source"), doc.metadata.get("chunk_id"))
                if key in seen_keys:
                    continue
                seen_keys.add(key)
                meta = dict(doc.metadata)
                if idx in dense_ranked:
                    meta["score"] = float(distances[0][dense_ranked.index(idx)])
                else:
                    meta["score"] = 0.0
                meta["retrieval_mode"] = retrieval_mode
                candidates.append(Document(content=doc.content, metadata=meta))

            # --- Rerank candidates with cross-encoder, then keep top-k ---
            if reranking_on and len(candidates) > k:
                loop = asyncio.get_event_loop()
                results = await loop.run_in_executor(
                    None, lambda: self.reranker.rerank(query, candidates, k)
                )
                for doc in results:
                    doc.metadata["retrieval_mode"] = retrieval_mode + "+rerank"
            else:
                results = candidates[:k]

            return results

        except Exception as e:
            logger.error(f"Error retrieving context: {e}")
            return []

    # ------------------
    # Index cache helpers
    # ------------------
    def _cache_dir(self) -> str:
        return os.path.join(self.documents_path, ".index")

    def _manifest_path(self) -> str:
        return os.path.join(self._cache_dir(), "manifest.json")

    def _faiss_path(self) -> str:
        return os.path.join(self._cache_dir(), "index.faiss")

    def _chunks_path(self) -> str:
        return os.path.join(self._cache_dir(), "chunks.jsonl")

    def _build_manifest(self) -> Dict[str, Any]:
        files_info: List[Dict[str, Any]] = []
        for root, _, files in os.walk(self.documents_path):
            for filename in files:
                ext = Path(filename).suffix.lower()
                if ext in SUPPORTED_FILE_EXTENSIONS:
                    fp = os.path.join(root, filename)
                    try:
                        st = os.stat(fp)
                        files_info.append({
                            "path": os.path.relpath(fp, self.documents_path),
                            "size": st.st_size,
                            "mtime": int(st.st_mtime),
                        })
                    except Exception:
                        continue
        return {"files": sorted(files_info, key=lambda x: x["path"]) }

    def _load_index_cache(self) -> bool:
        manifest_path = self._manifest_path()
        faiss_path = self._faiss_path()
        chunks_path = self._chunks_path()
        if not (os.path.exists(manifest_path) and os.path.exists(faiss_path) and os.path.exists(chunks_path)):
            return False
        # Compare manifests
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                saved_manifest = json.load(f)
        except Exception:
            return False
        current_manifest = self._build_manifest()
        if saved_manifest != current_manifest:
            return False
        # Load chunks
        documents: List[Document] = []
        try:
            with open(chunks_path, "r", encoding="utf-8") as f:
                for line in f:
                    if not line.strip():
                        continue
                    rec = json.loads(line)
                    documents.append(Document(content=rec["content"], metadata=rec["metadata"]))
        except Exception:
            return False
        # Load index
        try:
            index = faiss.read_index(faiss_path)
        except Exception:
            return False
        # Assign
        self.documents = documents
        self.index = index
        # embeddings not needed at query time for FAISS
        self.embeddings = None
        return True

    def _save_index_cache(self) -> None:
        cache_dir = self._cache_dir()
        os.makedirs(cache_dir, exist_ok=True)
        # Save manifest
        manifest = self._build_manifest()
        with open(self._manifest_path(), "w", encoding="utf-8") as f:
            json.dump(manifest, f)
        # Save chunks
        with open(self._chunks_path(), "w", encoding="utf-8") as f:
            for doc in self.documents:
                f.write(json.dumps({"content": doc.content, "metadata": doc.metadata}, ensure_ascii=False) + "\n")
        # Save FAISS index
        if self.index is not None:
            faiss.write_index(self.index, self._faiss_path())

class PlannerAgent(Agent):
    """Agent that creates research plans"""
    
    def __init__(self, llm_provider: LLMProvider):
        super().__init__("PlannerAgent", llm_provider)
    
    async def process(self, research_query: ResearchQuery) -> Dict[str, Any]:
        """Create a research plan based on the query"""
        
        prompt = f"""Create a research plan for this query:
Query: "{research_query.query}"
Intent: {research_query.intent}
Sub-queries: {', '.join(research_query.sub_queries)}

Provide a JSON research plan:
{{
    "steps": ["step 1", "step 2", "step 3"],
    "focus_areas": ["area 1", "area 2"],
    "output_structure": ["Introduction", "Main Findings", "Conclusion"],
    "complexity": "medium"
}}

Plan:"""
        
        response = await self.llm_provider.generate_text(prompt, max_tokens=300)
        
        try:
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response[json_start:json_end]
                return json.loads(json_str)
        except:
            pass
        
        # Fallback plan
        return {
            "steps": ["Gather information", "Analyze findings", "Synthesize results"],
            "focus_areas": [research_query.intent.split()[-1] if research_query.intent else "research"],
            "output_structure": ["Introduction", "Key Findings", "Analysis", "Conclusion"],
            "complexity": "medium"
        }

class SearcherAgent(Agent):
    """Agent that searches and retrieves relevant information"""

    def __init__(
        self,
        llm_provider: LLMProvider,
        rag_retriever: RAGRetriever,
        use_hyde: bool = True,
        web_searcher: Optional["WebSearcher"] = None,
    ):
        super().__init__("SearcherAgent", llm_provider)
        self.rag_retriever = rag_retriever
        self.use_hyde = use_hyde
        self.web_searcher = web_searcher

    async def _generate_hyde_passage(self, query: str) -> Optional[str]:
        """Generate a hypothetical document passage for the query (HyDE).

        The generated passage is embedded instead of the raw query for dense retrieval.
        It resembles a real document answer, so it aligns better with the corpus embeddings.
        """
        prompt = (
            f'Write a short, factual paragraph (3-5 sentences) that directly answers '
            f'the following question as if from a research document:\n\n'
            f'Question: "{query}"\n\nParagraph:'
        )
        try:
            passage = await self.llm_provider.generate_text(prompt, max_tokens=150)
            passage = passage.strip()
            return passage if len(passage) > 20 else None
        except Exception as e:
            logger.warning(f"HyDE generation failed: {e}")
            return None

    async def process(self, query: str) -> Dict[str, Any]:
        """Search for relevant information using RAG + web search in parallel."""

        # HyDE: generate a hypothetical passage for dense retrieval
        hyde_passage: Optional[str] = None
        if self.use_hyde:
            hyde_passage = await self._generate_hyde_passage(query)

        # Fire RAG and web search in parallel
        rag_task = self.rag_retriever.retrieve_context(query, k=3, hyde_passage=hyde_passage)
        web_task = (
            self.web_searcher.search(query, max_results=3)
            if self.web_searcher and self.web_searcher.is_available
            else asyncio.sleep(0, result=[])
        )
        relevant_docs, web_docs = await asyncio.gather(rag_task, web_task)

        all_docs: List[Document] = list(relevant_docs) + list(web_docs)

        if not all_docs:
            return {
                "content": "No relevant documents found in the knowledge base or web.",
                "sources": [],
                "relevance_score": 0.0,
                "web_results": 0,
            }

        # Build combined content — label web results so LLM knows their origin
        content_pieces, sources, scored = [], [], []
        web_count = 0
        for doc in all_docs:
            is_web = doc.metadata.get("type") == "web"
            content_pieces.append(f"[Web] {doc.content}" if is_web else doc.content)
            sources.append(doc.metadata.get("source", "unknown"))
            scored.append(doc.metadata.get("score"))
            if is_web:
                web_count += 1

        combined_content = "\n\n".join(content_pieces)

        # Use LLM to process the content
        prompt = f"""Based on the following retrieved information, provide a focused summary for this query: "{query}"

Retrieved Information:
{combined_content}

Please provide a clear, organized summary that directly addresses the query:"""

        processed_content = await self.llm_provider.generate_text(prompt, max_tokens=400)

        # Estimate relevance
        valid_scores = [s for s in scored if s is not None]
        if valid_scores:
            relevance = float(np.mean(valid_scores))
            relevance_norm = max(0.0, min((relevance + 1.0) / 2.0, 1.0))
        else:
            relevance_norm = min(len(all_docs) / float(DEFAULT_TOP_K), 1.0)

        return {
            "content": processed_content,
            "sources": list(set(sources)),
            "relevance_score": relevance_norm,
            "web_results": web_count,
        }

class SynthesizerAgent(Agent):
    """Agent that synthesizes information into comprehensive reports"""
    
    def __init__(self, llm_provider: LLMProvider):
        super().__init__("SynthesizerAgent", llm_provider)
    
    async def process(self, research_data: Dict[str, Any]) -> ResearchResult:
        """Synthesize research findings into a comprehensive report"""
        
        query = research_data.get("query", "")
        search_results = research_data.get("search_results", [])
        plan = research_data.get("plan", {})
        
        # Combine all search results
        all_content = []
        all_sources = set()
        avg_confidence = 0.0
        
        for result in search_results:
            all_content.append(result.get("content", ""))
            all_sources.update(result.get("sources", []))
            avg_confidence += result.get("relevance_score", 0.0)
        
        if search_results:
            avg_confidence /= len(search_results)
        
        combined_content = "\n\n---\n\n".join(all_content)
        
        # Include prior conversation context so the report builds on previous answers
        prior_context = research_data.get("prior_context", "")
        context_block = (
            f"\n\nPrior conversation context (do not repeat, but build upon it):\n{prior_context}\n"
            if prior_context else ""
        )

        prompt = f"""Create a comprehensive research report for: "{query}"

Research Plan: {json.dumps(plan, indent=2)}
{context_block}
Gathered Information:
{combined_content}

Please create a well-structured research report with clear sections, key insights, and conclusions. Format it professionally with headings and organized content:"""
        
        report = await self.llm_provider.generate_text(prompt, max_tokens=600)
        
        # Ensure the report has proper structure
        if not any(marker in report for marker in ["#", "**", "Introduction", "Conclusion"]):
            report = f"""# Research Report: {query}

## Introduction
This report addresses the research query: "{query}"

## Key Findings
{report}

## Conclusion
The research provides valuable insights into the topic and highlights important areas for further investigation.
"""
        
        # Append references if available
        if all_sources:
            references_section = "\n\n## References\n" + "\n".join(f"- {src}" for src in sorted(all_sources))
            if "## References" not in report:
                report = report.strip() + references_section
        
        return ResearchResult(
            content=report,
            sources=list(all_sources),
            confidence=avg_confidence,
            agent_chain=["QueryUnderstandingAgent", "PlannerAgent", "SearcherAgent", "SynthesizerAgent"]
        )

class AgenticResearchAssistant:
    """Main orchestrator for the agentic research system"""
    
    def __init__(self, llm_provider: str = "huggingface", documents_path: str = "./documents", model_name: Optional[str] = None, architecture: str = "native"):
        # Initialize LLM provider (no API key needed!)
        self.llm_provider = LLMProvider(llm_provider, model_name)
        
        # Initialize RAG retriever
        self.rag_retriever = RAGRetriever(documents_path)
        
        # Initialize agents
        self.query_agent = QueryUnderstandingAgent(self.llm_provider)
        self.planner_agent = PlannerAgent(self.llm_provider)
        self.web_searcher = WebSearcher()
        self.searcher_agent = SearcherAgent(
            self.llm_provider, self.rag_retriever,
            use_hyde=True, web_searcher=self.web_searcher,
        )
        self.synthesizer_agent = SynthesizerAgent(self.llm_provider)
        
        # Architecture selection: 'native' or 'langgraph'
        self.architecture = architecture if architecture in ("native", "langgraph") else "native"
        
        # LangChain/LangGraph pipeline (required deps available)
        self.lc_enabled = self.architecture == "langgraph"
        if self.lc_enabled:
            self._initialize_langchain_components(documents_path)
        
        logger.info("Agentic Research Assistant initialized")
        if self.lc_enabled:
            logger.info("LangChain/LangGraph mode enabled")
    
    async def research(
        self,
        query: str,
        memory: Optional["ConversationMemory"] = None,
    ) -> ResearchResult:
        """Main research method that orchestrates all agents.

        Args:
            query:  The user's question (may be a follow-up referencing prior turns).
            memory: Optional conversation memory for query rewriting and synthesis context.
        """
        logger.info(f"🔍 Starting research for: '{query}'")

        # If LangGraph mode is enabled, run the graph orchestration
        if self.lc_enabled:
            try:
                return await self._research_with_langgraph(query)
            except Exception as e:
                logger.error(f"LangGraph path failed: {e}. Falling back to native.")

        try:
            # Step 1: Understand the query (with optional rewrite for follow-ups)
            logger.info("Step 1: Analyzing query...")
            if memory and not memory.is_empty():
                query = await self.query_agent.rewrite_query(query, memory)
                logger.info(f"✅ Query rewritten to: '{query}'")
            research_query = await self.query_agent.process(query)
            logger.info(f"✅ Query analyzed. Intent: {research_query.intent}")
            
            # Step 2: Create research plan
            logger.info("Step 2: Creating research plan...")
            plan = await self.planner_agent.process(research_query)
            logger.info(f"✅ Research plan created")
            
            # Step 3: Search for information
            logger.info("Step 3: Searching for relevant information...")
            
            async def search_one(sub_q: str) -> Dict[str, Any]:
                logger.info(f"  🔎 Searching: '{sub_q}'")
                return await self.searcher_agent.process(sub_q)

            # Run sub-query searches in parallel for speed
            tasks = [search_one(sub_q) for sub_q in research_query.sub_queries]
            search_results = await asyncio.gather(*tasks)
            
            logger.info(f"✅ Search complete ({len(search_results)} results)")
            
            # Step 4: Synthesize results
            logger.info("Step 4: Synthesizing research findings...")
            research_data = {
                "query": query,
                "search_results": search_results,
                "plan": plan,
                "research_query": research_query,
                "prior_context": memory.format_for_synthesis() if memory and not memory.is_empty() else "",
            }
            
            final_result = await self.synthesizer_agent.process(research_data)
            logger.info("✅ Research synthesis complete!")
            
            return final_result
            
        except Exception as e:
            logger.error(f"❌ Error during research: {e}")
            return ResearchResult(
                content=f"An error occurred during research: {str(e)}",
                sources=[],
                confidence=0.0,
                agent_chain=["Error"]
            )

    # ==============================
    # LangChain / LangGraph pipeline
    # ==============================
    def _initialize_langchain_components(self, documents_path: str):
        """Initialize LangChain embeddings, vector store, and LLM for LangGraph mode."""
        try:
            # Embeddings
            self.lc_embeddings = LCHuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

            # Load and chunk documents (reuse native loader for consistency)
            docs: List[Document] = []
            if not os.path.exists(documents_path):
                os.makedirs(documents_path, exist_ok=True)
            for filename in os.listdir(documents_path):
                if filename.endswith(".txt"):
                    path = os.path.join(documents_path, filename)
                    try:
                        with open(path, "r", encoding="utf-8") as f:
                            content = f.read().strip()
                        if content:
                            for i, chunk in enumerate(self.rag_retriever._split_text(content)):
                                docs.append(Document(content=chunk, metadata={"source": filename, "chunk_id": i}))
                    except Exception as e:
                        logger.error(f"Error reading {filename} for LangChain index: {e}")

            texts = [d.content for d in docs]
            metadatas = [d.metadata for d in docs]
            # Vector store
            self.lc_vectorstore = LCFAISS.from_texts(texts=texts, embedding=self.lc_embeddings, metadatas=metadatas)
            self.lc_retriever = self.lc_vectorstore.as_retriever(search_kwargs={"k": 3})

            # LLM
            if self.llm_provider.provider_type == "ollama":
                # Use LangChain Ollama wrapper
                self.lc_llm = LCOllama(
                    model=getattr(self.llm_provider, "ollama_model_name", os.getenv("OLLAMA_MODEL", "mistral:7b-instruct-q5_K_M")),
                    num_ctx=getattr(self.llm_provider, "ollama_num_ctx", int(os.getenv("OLLAMA_NUM_CTX", "8192"))),
                    temperature=getattr(self.llm_provider, "ollama_temperature", float(os.getenv("OLLAMA_TEMPERATURE", "0.3"))),
                )
            else:
                # Use the Hugging Face pipeline through LangChain
                # Reuse the provider's pipeline if available; else create a small default
                if isinstance(self.llm_provider.pipeline, type(None)):
                    hf_pipe = pipeline(
                        "text-generation",
                        model=os.getenv("HF_MODEL_NAME", "distilgpt2"),
                        device=0 if torch.cuda.is_available() else -1,
                    )
                else:
                    hf_pipe = self.llm_provider.pipeline
                self.lc_llm = LCHuggingFacePipeline(pipeline=hf_pipe)

            # Build LangGraph
            self.lg_compiled = self._build_langgraph()
        except Exception as e:
            logger.error(f"Failed to initialize LangChain/LangGraph: {e}")
            self.lc_enabled = False

    def _build_langgraph(self):
        """Create and compile a LangGraph that mirrors the native agent flow."""
        # State is a dict with these keys
        # { query: str, analysis: dict, plan: dict, sub_queries: List[str], sub_results: List[dict], report: str, sources: List[str] }

        def analyze_node(state: Dict[str, Any]) -> Dict[str, Any]:
            query = state["query"]
            prompt = f"""Analyze this research query and provide JSON output:\nQuery: \"{query}\"\n\nProvide analysis in JSON format:\n{{\n    \"intent\": \"brief description of what user wants\",\n    \"sub_queries\": [\"focused sub-query 1\", \"focused sub-query 2\"],\n    \"priority\": 3,\n    \"research_domain\": \"domain like AI, technology, science, etc.\"\n}}\n\nJSON:"""
            try:
                response = self.lc_llm.invoke(prompt)
                json_start = response.find('{')
                json_end = response.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    analysis = json.loads(response[json_start:json_end])
                else:
                    analysis = {"intent": f"Research about {query}", "sub_queries": [query], "priority": 3}
            except Exception:
                analysis = {"intent": f"Research about {query}", "sub_queries": [query], "priority": 3}
            return {"analysis": analysis, "sub_queries": analysis.get("sub_queries", [query])[:3]}

        def plan_node(state: Dict[str, Any]) -> Dict[str, Any]:
            analysis = state.get("analysis", {})
            query = state["query"]
            prompt = f"""Create a research plan for this query:\nQuery: \"{query}\"\nIntent: {analysis.get('intent', '')}\nSub-queries: {', '.join(state.get('sub_queries', []))}\n\nProvide a JSON research plan:\n{{\n    \"steps\": [\"step 1\", \"step 2\", \"step 3\"],\n    \"focus_areas\": [\"area 1\", \"area 2\"],\n    \"output_structure\": [\"Introduction\", \"Main Findings\", \"Conclusion\"],\n    \"complexity\": \"medium\"\n}}\n\nPlan:"""
            try:
                response = self.lc_llm.invoke(prompt)
                json_start = response.find('{')
                json_end = response.rfind('}') + 1
                plan = json.loads(response[json_start:json_end]) if json_start >= 0 and json_end > json_start else {}
            except Exception:
                plan = {}
            if not plan:
                plan = {"steps": ["Gather information", "Analyze findings", "Synthesize results"],
                        "focus_areas": [analysis.get("research_domain", "research")],
                        "output_structure": ["Introduction", "Key Findings", "Analysis", "Conclusion"],
                        "complexity": "medium"}
            return {"plan": plan}

        def search_node(state: Dict[str, Any]) -> Dict[str, Any]:
            sub_queries = state.get("sub_queries", [])
            sub_results: List[Dict[str, Any]] = []
            all_sources = set()
            for sq in sub_queries:
                # Retrieve LangChain docs
                try:
                    lc_docs = self.lc_retriever.get_relevant_documents(sq)
                except Exception as e:
                    logger.error(f"LangChain retrieval error: {e}")
                    lc_docs = []
                sources = []
                combined = []
                for d in lc_docs:
                    # d.page_content, d.metadata
                    combined.append(d.page_content)
                    src = d.metadata.get("source", "unknown") if hasattr(d, "metadata") else "unknown"
                    sources.append(src)
                combined_text = "\n\n".join(combined) if combined else ""
                prompt = f"""Based on the following retrieved information, provide a focused summary for this query: \"{sq}\"\n\nRetrieved Information:\n{combined_text}\n\nPlease provide a clear, organized summary that directly addresses the query:"""
                try:
                    summary = self.lc_llm.invoke(prompt)
                except Exception:
                    summary = combined_text[:800]
                sub_results.append({
                    "content": summary,
                    "sources": list(set(sources)),
                    "relevance_score": 1.0 if combined_text else 0.0,
                })
                all_sources.update(sources)
            return {"sub_results": sub_results, "sources": list(all_sources)}

        def synthesize_node(state: Dict[str, Any]) -> Dict[str, Any]:
            query = state["query"]
            plan = state.get("plan", {})
            sub_results = state.get("sub_results", [])
            combined_content = "\n\n---\n\n".join([r.get("content", "") for r in sub_results])
            prompt = f"""Create a comprehensive research report for: \"{query}\"\n\nResearch Plan: {json.dumps(plan, indent=2)}\n\nGathered Information:\n{combined_content}\n\nPlease create a well-structured research report with clear sections, key insights, and conclusions. Format it professionally with headings and organized content:"""
            try:
                report = self.lc_llm.invoke(prompt)
            except Exception:
                report = combined_content
            # Ensure minimal structure
            if not any(marker in report for marker in ["#", "**", "Introduction", "Conclusion"]):
                report = f"# Research Report: {query}\n\n## Key Findings\n{report}\n\n## Conclusion\nThis report summarizes retrieved context."
            return {"report": report}

        graph = StateGraph(dict)
        graph.add_node("analyze", analyze_node)
        graph.add_node("plan", plan_node)
        graph.add_node("search", search_node)
        graph.add_node("synthesize", synthesize_node)

        graph.set_entry_point("analyze")
        graph.add_edge("analyze", "plan")
        graph.add_edge("plan", "search")
        graph.add_edge("search", "synthesize")
        graph.add_edge("synthesize", END)

        return graph.compile()

    async def _research_with_langgraph(self, query: str) -> ResearchResult:
        """Run the LangGraph pipeline and return a ResearchResult."""
        initial_state = {"query": query}
        # Run synchronously since LangGraph invoke is sync; wrap in thread to avoid blocking loop
        loop = asyncio.get_event_loop()
        state = await loop.run_in_executor(None, lambda: self.lg_compiled.invoke(initial_state))

        sub_results = state.get("sub_results", [])
        avg_conf = sum(r.get("relevance_score", 0.0) for r in sub_results) / (len(sub_results) or 1)
        sources = set()
        for r in sub_results:
            sources.update(r.get("sources", []))

        report = state.get("report", "")
        if sources and "## References" not in report:
            report = report.rstrip() + "\n\n## References\n" + "\n".join(f"- {s}" for s in sorted(sources))

        return ResearchResult(
            content=report,
            sources=list(sorted(sources)),
            confidence=avg_conf,
            agent_chain=["LangGraph: analyze", "LangGraph: plan", "LangGraph: search", "LangGraph: synthesize"],
        )
    
    def add_documents(self, documents: List[str], filenames: List[str] = None):
        """Add new documents to the knowledge base"""
        if filenames is None:
            filenames = [f"doc_{i}.txt" for i in range(len(documents))]
        
        for doc, filename in zip(documents, filenames):
            filepath = os.path.join(self.rag_retriever.documents_path, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(doc)
        
        # Rebuild index
        self.rag_retriever._load_and_index_documents()
        logger.info(f"✅ Added {len(documents)} documents to knowledge base")

    def download_arxiv_abstracts(self, query: str, max_results: int = 50) -> int:
        """Download arXiv abstracts for a query and add them as documents.
        Returns number of files added.
        """
        try:
            import urllib.parse
            import xml.etree.ElementTree as ET
            base = "http://export.arxiv.org/api/query"
            params = {
                "search_query": f"all:{query}",
                "start": 0,
                "max_results": max(1, min(max_results, 200)),
                "sortBy": "submittedDate",
                "sortOrder": "descending",
            }
            url = base + "?" + urllib.parse.urlencode(params)
            resp = requests.get(url, timeout=30)
            if resp.status_code != 200:
                logger.error(f"arXiv API error: HTTP {resp.status_code}")
                return 0
            # Parse Atom XML
            try:
                root = ET.fromstring(resp.text)
            except Exception as e:
                logger.error(f"arXiv XML parse error: {e}")
                return 0
            ns = {"atom": "http://www.w3.org/2005/Atom"}
            entries = root.findall("atom:entry", ns)
            if not entries:
                return 0
            added = 0
            for i, entry in enumerate(entries):
                title_el = entry.find("atom:title", ns)
                summary_el = entry.find("atom:summary", ns)
                title = (title_el.text or "Untitled").strip()
                summary = (summary_el.text or "").strip()
                # Build file content
                content = f"Title: {title}\n\nAbstract:\n{summary}\n"
                # Sanitize filename
                safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "-", "_"))
                safe_title = "_".join(safe_title.split())[:80]
                filename = f"arxiv_{safe_title or 'paper'}_{i+1}.txt"
                filepath = os.path.join(self.rag_retriever.documents_path, filename)
                try:
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(content)
                    added += 1
                except Exception as e:
                    logger.error(f"Write error {filename}: {e}")
            if added:
                self.rag_retriever._load_and_index_documents()
            return added
        except Exception as e:
            logger.error(f"arXiv download error: {e}")
            return 0

# CLI Interface and Usage Examples
async def main():
    """Interactive CLI for the Research Assistant"""
    
    print("\n🤖 Agentic RAG-Powered AI Research Assistant")
    print("🆓 Open Source Mode - No API Keys Required!")
    print("=" * 55)
    
    print("\n📋 Available LLM Options:")
    print("1. HuggingFace Transformers (Local)")
    print("2. Ollama (Local - requires installation)")
    print("3. Fallback Mode (Rule-based)")
    print("4. Google Gemini (API)")
    print("5. Groq (API)")
    print("6. OpenAI (API)")
    
    print("\n🧠 Orchestration Architecture:")
    print("a. Native agents (default)")
    print("b. LangGraph (requires langchain & langgraph)")
    
    while True:
        try:
            choice = input("\nSelect LLM option (1-3) or press Enter for default: ").strip()
            if not choice:
                choice = "1"
            
            if choice == "1":
                llm_provider = "huggingface"
                llm_model_name = os.getenv("HF_MODEL_NAME", "distilgpt2")
                break
            elif choice == "2":
                llm_provider = "ollama"
                print("📝 Note: Make sure Ollama is running with: ollama serve")
                # Ask for model tag or use env/default
                entered = input("Enter Ollama model tag (Enter for default '" + os.getenv("OLLAMA_MODEL", "mistral:7b-instruct-q5_K_M") + "'): ").strip()
                llm_model_name = entered if entered else os.getenv("OLLAMA_MODEL", "mistral:7b-instruct-q5_K_M")
                break
            elif choice == "3":
                llm_provider = "fallback"
                llm_model_name = None
                break
            elif choice == "4":
                llm_provider = "gemini"
                llm_model_name = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
                print("🔐 Requires GOOGLE_API_KEY env var")
                break
            elif choice == "5":
                llm_provider = "groq"
                llm_model_name = os.getenv("GROQ_MODEL", "llama3-8b-8192")
                print("🔐 Requires GROQ_API_KEY env var")
                break
            elif choice == "6":
                llm_provider = "openai"
                llm_model_name = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
                print("🔐 Requires OPENAI_API_KEY env var")
                break
            else:
                print("Please select 1-6")
        except KeyboardInterrupt:
            print("\nExiting...")
            return
    # Architecture choice
    arch_choice = input("\nSelect architecture (a/b) or press Enter for default (a): ").strip().lower() or "a"
    architecture = "langgraph" if arch_choice == "b" else "native"
    
    # Initialize assistant
    print(f"\n🚀 Initializing with {llm_provider} backend...")
    assistant = AgenticResearchAssistant(llm_provider, model_name=llm_model_name, architecture=architecture)
    
    # Example queries
    example_queries = [
        "latest trends in AI safety",
        "machine learning applications in healthcare",
        "natural language processing advancements",
        "deep learning fundamentals",
        "research methodology in AI"
    ]
    
    while True:
        print("\n" + "="*50)
        print("🎯 What would you like to research?")
        print("\nOptions:")
        print("1. 📝 Enter custom research query")
        print("2. 📚 Try example queries")
        print("3. 📄 Add documents to knowledge base")
        print("4. 📦 Bulk import a folder (recursive)")
        print("5. ⬇️  Enrich KB: Download arXiv abstracts for a topic")
        print("6. 🔍 View current documents")
        print("7. ❌ Exit")
        
        try:
            choice = input("\nSelect option (1-7): ").strip()
            
            if choice == "1":
                query = input("\n📝 Enter your research query: ").strip()
                if query:
                    print(f"\n🔍 Researching: '{query}'")
                    print("-" * 50)
                    
                    start_time = time.time()
                    result = await assistant.research(query)
                    elapsed_time = time.time() - start_time
                    
                    print("\n" + "="*50)
                    print("📋 RESEARCH REPORT")
                    print("="*50)
                    print(result.content)
                    
                    if result.sources:
                        print(f"\n📚 Sources: {', '.join(result.sources)}")
                    
                    print(f"\n📊 Confidence: {result.confidence:.2f}")
                    print(f"🔄 Agent Chain: {' → '.join(result.agent_chain)}")
                    print(f"⏱️  Processing Time: {elapsed_time:.1f}s")
            
            elif choice == "2":
                print("\n📚 Example Queries:")
                for i, query in enumerate(example_queries, 1):
                    print(f"{i}. {query}")
                
                try:
                    query_choice = int(input("\nSelect query (1-5): ")) - 1
                    if 0 <= query_choice < len(example_queries):
                        query = example_queries[query_choice]
                        print(f"\n🔍 Researching: '{query}'")
                        print("-" * 50)
                        
                        start_time = time.time()
                        result = await assistant.research(query)
                        elapsed_time = time.time() - start_time
                        
                        print("\n" + "="*50)
                        print("📋 RESEARCH REPORT")
                        print("="*50)
                        print(result.content)
                        
                        if result.sources:
                            print(f"\n📚 Sources: {', '.join(result.sources)}")
                        
                        print(f"\n📊 Confidence: {result.confidence:.2f}")
                        print(f"🔄 Agent Chain: {' → '.join(result.agent_chain)}")
                        print(f"⏱️  Processing Time: {elapsed_time:.1f}s")
                    else:
                        print("❌ Invalid selection")
                except ValueError:
                    print("❌ Please enter a valid number")
            
            elif choice == "3":
                print("\n📄 Add Document to Knowledge Base")
                print("Enter document content (type 'END' on a new line to finish):")
                
                content_lines = []
                while True:
                    try:
                        line = input()
                        if line.strip().upper() == "END":
                            break
                        content_lines.append(line)
                    except KeyboardInterrupt:
                        print("\n❌ Document addition cancelled")
                        break
                
                if content_lines:
                    content = "\n".join(content_lines)
                    filename = input("\n📝 Enter filename (or press Enter for auto-name): ").strip()
                    
                    if not filename:
                        doc_count = len([f for f in os.listdir(assistant.rag_retriever.documents_path) if f.endswith('.txt')])
                        filename = f"user_document_{doc_count + 1}.txt"
                    
                    if not filename.endswith('.txt'):
                        filename += '.txt'
                    
                    assistant.add_documents([content], [filename])
                    print(f"✅ Document added as '{filename}'")
                else:
                    print("❌ No content provided")
            
            elif choice == "4":
                print("\n📦 Bulk Import")
                folder = input("Enter folder path to import recursively: ").strip()
                if folder:
                    try:
                        imported = 0
                        # Copy files preserving structure under documents_path
                        src_path = Path(folder)
                        dst_root = Path(assistant.rag_retriever.documents_path)
                        for root, _, files in os.walk(src_path):
                            for filename in files:
                                ext = Path(filename).suffix.lower()
                                if ext in {'.txt', '.md', '.pdf', '.docx', '.html', '.htm'}:
                                    src_file = Path(root) / filename
                                    rel = os.path.relpath(src_file, src_path)
                                    dst_file = dst_root / rel
                                    dst_file.parent.mkdir(parents=True, exist_ok=True)
                                    try:
                                        shutil.copy2(src_file, dst_file)
                                        imported += 1
                                    except Exception as e:
                                        print(f"Skip {src_file}: {e}")
                        # Re-index
                        assistant.rag_retriever._load_and_index_documents()
                        print(f"✅ Imported {imported} files and rebuilt index")
                    except Exception as e:
                        print(f"❌ Import error: {e}")
                else:
                    print("❌ No folder provided")

            elif choice == "5":
                topic = input("Enter arXiv topic/keywords (e.g., 'AI safety'): ").strip()
                try:
                    max_results = int(input("Max results to fetch (default 25): ") or "25")
                except ValueError:
                    max_results = 25
                if topic:
                    print(f"Fetching up to {max_results} arXiv abstracts for: {topic}")
                    added = assistant.download_arxiv_abstracts(topic, max_results=max_results)
                    print(f"✅ Added {added} abstracts to the knowledge base")
                else:
                    print("❌ No topic provided")

            elif choice == "6":
                print("\n📚 Current Documents in Knowledge Base:")
                try:
                    docs_path = assistant.rag_retriever.documents_path
                    if os.path.exists(docs_path):
                        # Walk recursively and list a sample of supported files
                        supported = []
                        for root, _, files in os.walk(docs_path):
                            for filename in files:
                                if Path(filename).suffix.lower() in {'.txt', '.md', '.pdf', '.docx', '.html', '.htm'}:
                                    supported.append(os.path.join(root, filename))
                        if supported:
                            for i, filepath in enumerate(sorted(supported)[:50], 1):
                                rel = os.path.relpath(filepath, docs_path)
                                preview = ""
                                try:
                                    text = assistant.rag_retriever._read_document_file(filepath)
                                    preview = (text[:100] + "...") if len(text) > 100 else text
                                    word_count = len(text.split())
                                except Exception as e:
                                    preview = f"(Error reading file: {e})"
                                    word_count = 0
                                print(f"{i}. {rel} ({word_count} words)")
                                if preview:
                                    print(f"   Preview: {preview}")
                        else:
                            print("📭 No documents found")
                    else:
                        print("📭 Documents directory not found")
                except Exception as e:
                    print(f"❌ Error listing documents: {e}")
            
            elif choice == "7":
                print("\n👋 Thank you for using the Agentic Research Assistant!")
                print("🌟 Built with open-source technologies - no API keys needed!")
                break
            
            else:
                print("❌ Invalid option. Please select 1-5.")
                
        except KeyboardInterrupt:
            print("\n\n👋 Goodbye!")
            break
        except Exception as e:
            print(f"❌ An error occurred: {e}")

def print_installation_guide():
    """Print installation instructions"""
    print("\n" + "="*60)
    print("📦 INSTALLATION GUIDE")
    print("="*60)
    print("\n1️⃣  Install Python Dependencies:")
    print("   pip install sentence-transformers faiss-cpu transformers torch numpy")
    
    print("\n2️⃣  Optional: Install Ollama for Better LLM Support")
    print("   - Download from: https://ollama.ai")
    print("   - Install a model: ollama pull llama2")
    print("   - Start server: ollama serve")
    
    print("\n3️⃣  Run the Assistant:")
    print("   python agentic_rag_assistant.py")
    
    print("\n🎯 FEATURES:")
    print("✅ No API keys required")
    print("✅ Runs completely offline")
    print("✅ Multi-agent RAG system") 
    print("✅ Vector similarity search")
    print("✅ Extensible document knowledge base")
    print("✅ Interactive CLI interface")
    
    print("\n🔧 SUPPORTED LLM BACKENDS:")
    print("• HuggingFace Transformers (Local)")
    print("• Ollama (Local installation)")
    print("• Rule-based fallback (Always works)")

if __name__ == "__main__":
    print("🤖 Agentic RAG-Powered AI Research Assistant")
    print("🆓 Open Source Edition - No API Keys Required!")
    print("=" * 55)
    
    # Check if required packages are installed
    missing_packages = []
    
    # Load .env if available
    if isinstance(load_dotenv, Callable):
        try:
            load_dotenv()
            logger.info("Loaded environment variables from .env (if present)")
        except Exception:
            pass

    # Probe key packages lazily so IDEs don't flag unused imports
    for mod, pip_name in [
        ("sentence_transformers", "sentence-transformers"),
        ("faiss", "faiss-cpu"),
        ("transformers", "transformers"),
        ("torch", "torch"),
        ("langchain", "langchain"),
        ("langgraph", "langgraph"),
        ("langchain_community", "langchain-community"),
    ]:
        try:
            importlib.import_module(mod)
        except Exception:
            missing_packages.append(pip_name)
    
    if missing_packages:
        print(f"\n❌ Missing required packages: {', '.join(missing_packages)}")
        print("\n📦 Install with:")
        print(f"   pip install {' '.join(missing_packages)}")
        print_installation_guide()
        print("\nPlease install the required packages and run again.")
    else:
        print("✅ All dependencies found!")
        print("\n🚀 Starting Assistant...")
        
        # Run the assistant
        try:
            asyncio.run(main())
        except KeyboardInterrupt:
            print("\n👋 Assistant terminated by user")
        except Exception as e:
            print(f"\n❌ Fatal error: {e}")
            print("Please check the installation and try again.")